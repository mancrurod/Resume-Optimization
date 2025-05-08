import os
import sys
import logging
import webbrowser
import docx 
from datetime import datetime
from dotenv import load_dotenv, find_dotenv

from src.convert_to_md import convert_docx_to_md
from src.optimize_resume import generate_prompt
from src.adapt_resume import adapt_resume
from src.export_resume import convert_md_to_html, convert_html_to_pdf, edit_html_content

def setup_logger() -> logging.Logger:
    """
    Create and configure a logger that logs to a timestamped file inside logs/.
    
    Returns:
        Logger: Configured logger instance.
    """
    os.makedirs("logs", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_filename = f"logs/resume_{timestamp}.log"

    logger = logging.getLogger("resume_logger")
    logger.setLevel(logging.INFO)

    file_handler = logging.FileHandler(log_filename, encoding="utf-8")
    formatter = logging.Formatter("[%(asctime)s] %(levelname)s - %(message)s")
    file_handler.setFormatter(formatter)

    logger.addHandler(file_handler)

    # Optional: also log to console
    stream_handler = logging.StreamHandler(sys.stdout)
    stream_handler.setFormatter(formatter)
    logger.addHandler(stream_handler)

    logger.info("🚀 Resume Optimization Pipeline started.")
    return logger


def validate_api_keys(logger) -> None:
    """
    Ensure that all required API keys are available in the environment.
    Exits the program if any key is missing.
    """
    missing = []
    if not os.getenv("OPENAI_API_KEY"):
        missing.append("OPENAI_API_KEY")
    if not os.getenv("GOOGLE_API_KEY"):
        missing.append("GOOGLE_API_KEY")
    if missing:
        logger.info(f"❌ Missing keys: {', '.join(missing)}. Please add them to your .env file.")
        sys.exit(1)
    logger.info("✅ API keys loaded successfully.")


def ensure_directories(dirs: list) -> None:
    """
    Create required directories if they don't already exist.

    Parameters:
        dirs (list): List of directory paths to create.
    """
    for d in dirs:
        os.makedirs(d, exist_ok=True)


def get_latest_docx_file(folder: str, logger) -> str:
    """
    Get the most recently modified .docx file from a folder.

    Parameters:
        folder (str): Directory to search.

    Returns:
        str: Filename of the latest .docx file.

    Raises:
        SystemExit: If no .docx file is found.
    """
    docs = [f for f in os.listdir(folder) if f.endswith(".docx")]
    if not docs:
        logger.info(f"❌ No .docx file found in '{folder}'")
        sys.exit(1)
    return sorted(docs, key=lambda f: os.path.getmtime(os.path.join(folder, f)))[-1]


def validate_docx_content(docx_path: str, logger) -> None:
    """
    Validate that the .docx file contains meaningful content.
    
    Parameters:
        docx_path (str): Full path to the .docx file.
        logger: Logger instance.
    
    Raises:
        SystemExit: If the file is empty or invalid.
    """
    try:
        doc = docx.Document(docx_path)
        has_paragraphs = any(p.text.strip() for p in doc.paragraphs)
        has_tables = any(t.rows for t in doc.tables)
        if not has_paragraphs and not has_tables:
            logger.info(f"❌ The .docx file '{docx_path}' is empty or contains no valid content.")
            sys.exit(1)
        logger.info("✅ .docx file validated: content found.")
    except Exception as e:
        logger.error(f"❌ Failed to open or read .docx file '{docx_path}': {e}")
        sys.exit(1)


def read_file(path: str, logger) -> str:
    """
    Read the content of a text file.

    Parameters:
        path (str): Path to the file.

    Returns:
        str: File content as a string.
    """
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception as e:
        logger.info(f"❌ Error reading {path}: {e}")
        sys.exit(1)


def save_file(path: str, content: str, logger) -> None:
    """
    Write content to a text file.

    Parameters:
        path (str): Path to the file.
        content (str): Text to write.
    """
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        logger.info(f"❌ Error saving {path}: {e}")
        sys.exit(1)


def main():
    """
    Execute the full resume optimization pipeline:
    1. Convert .docx to .md
    2. Generate LLM prompt
    3. Adapt resume using OpenAI or Gemini
    4. Convert to HTML
    5. Open visual HTML editor
    6. Export to PDF
    """

    # Setup logger
    logger = setup_logger()

    if not find_dotenv():
        logger.info("⚠️ .env file not found.")
    load_dotenv()
    validate_api_keys(logger)

    # Set directories
    input_dir = "original_docx"
    output_dir = "processed_cv"
    pdf_dir = "pdf_cv"
    job_path = "job_description.txt"

    ensure_directories([input_dir, output_dir, pdf_dir])

    # Step 1: Convert DOCX to Markdown
    logger.info("\n🔍 Step 1: Converting .docx to Markdown...")
    docx_filename = get_latest_docx_file(input_dir, logger)
    docx_path = os.path.join(input_dir, docx_filename)
    validate_docx_content(docx_path, logger)
    md_path = os.path.join(output_dir, os.path.splitext(docx_filename)[0] + ".md")
    convert_docx_to_md(docx_path, md_path)

    # Step 2: Generate LLM prompt
    logger.info("\n🧠 Step 2: Generating prompt for LLM...")
    md_content = read_file(md_path, logger)
    job_description = read_file(job_path, logger)
    prompt = generate_prompt(md_content, job_description)
    save_file(os.path.join(output_dir, "prompt.txt"), prompt, logger)

    # Step 3: Adapt resume via LLM
    logger.info("\n🤖 Step 3: Adapting resume using LLM...")
    adapted_md_path = os.path.join(output_dir, "adapted_resume.md")
    adapt_resume(os.path.join(output_dir, "prompt.txt"), adapted_md_path)

    # Step 4: Generate editable HTML
    logger.info("\n🌐 Step 4: Generating editable HTML for visual editor...")
    html_path = os.path.join(output_dir, os.path.splitext(docx_filename)[0] + ".html")
    convert_md_to_html(adapted_md_path, html_path, for_editor=True)

    # Step 5: Launch visual HTML editor (with Georgia font)
    logger.info("\n✍️ Step 5: Opening visual HTML editor...")
    edit_html_content(html_path)

    # Step 6: Export to PDF (based on edited HTML)
    logger.info("\n📄 Step 6: Exporting final resume to PDF...")
    pdf_path = os.path.join(pdf_dir, os.path.splitext(docx_filename)[0] + ".pdf")
    convert_html_to_pdf(html_path, pdf_path)

    logger.info(f"\n✅ DONE: Resume PDF saved at: {pdf_path}")
    try:
        webbrowser.open(f"file://{os.path.abspath(pdf_path)}")
        logger.info("📂 PDF opened in default viewer.")
    except Exception as e:
        logger.warning(f"⚠️ Could not open PDF automatically: {e}")




if __name__ == "__main__":
    main()

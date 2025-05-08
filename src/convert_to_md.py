import os
import re
import docx
from lxml import etree
from pathlib import Path

# === Style to Markdown mapping ===

STYLE_TO_MD = {
    "contacto": lambda text: f"[[CONTACT]]{text}",            
    "heading 1": lambda text: f"# {text}",               
    "heading 2": lambda text: f"## {text}",               
    "bullet": lambda text: f"- {text.lstrip('-• ').strip()}",
    "normal": lambda text: text,
    "key relevance": lambda text: f"**{text}**",
    }

DEFAULT_MD_MAPPING = lambda text: text


# === Helpers ===

def normalize_spacing(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    return text.strip() + "\n"

def extract_hyperlinks(doc: docx.Document) -> dict:
    return {
        r.rId: r._target
        for r in doc.part.rels.values()
        if "hyperlink" in r.reltype
    }

def extract_runs(para: docx.text.paragraph.Paragraph, rels: dict) -> str:
    result = []
    root = etree.fromstring(para._element.xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    seen_texts = set()

    for elem in root.xpath(".//w:hyperlink | .//w:r", namespaces=ns):
        if etree.QName(elem).localname == "hyperlink":
            rel_id = elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rel_id in rels:
                url = rels[rel_id]
                text = "".join(t.text for t in elem.xpath(".//w:t", namespaces=ns) if t.text).strip()
                norm_text = re.sub(r"\s+", " ", text)
                if norm_text and norm_text not in seen_texts:
                    result.append(f"[{norm_text}]({url})")
                    seen_texts.add(norm_text)
        elif etree.QName(elem).localname == "r":
            if elem.getparent() is not None and etree.QName(elem.getparent()).localname == "hyperlink":
                continue
            t = elem.xpath(".//w:t", namespaces=ns)
            if t and t[0].text:
                text = t[0].text.strip()
                norm_text = re.sub(r"\s+", " ", text)
                if not norm_text or norm_text in seen_texts:
                    continue
                is_bold = elem.xpath(".//w:b", namespaces=ns)
                is_italic = elem.xpath(".//w:i", namespaces=ns)

                if is_bold and is_italic:
                    result.append(f"***{norm_text}***")
                elif is_bold:
                    result.append(f"**{norm_text}**")
                elif is_italic:
                    result.append(f"*{norm_text}*")
                else:
                    result.append(norm_text)
                seen_texts.add(norm_text)

    return " ".join(result)


def process_paragraph(para: docx.text.paragraph.Paragraph, rels: dict) -> str:
    text = extract_runs(para, rels).strip()
    if not text:
        return ""
    style = para.style.name.strip().lower()
    if style in STYLE_TO_MD:
        return STYLE_TO_MD[style](text)
    if text.startswith("- ") or "•" in text:
        return f"- {text.lstrip('-• ').strip()}"
    return text


def process_table(table: docx.table.Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) >= 2:
        header_sep = "| " + " | ".join("---" for _ in rows[0].split("|") if _ and _ != " ") + " |"
        return "\n".join([rows[0], header_sep] + rows[1:])
    return "\n".join(rows)


def merge_education_blocks(md_lines: list) -> list:
    merged = []
    i = 0
    while i < len(md_lines):
        line = md_lines[i].strip()
        next_line = md_lines[i + 1].strip() if i + 1 < len(md_lines) else ""
        next_next_line = md_lines[i + 2].strip() if i + 2 < len(md_lines) else ""

        is_heading = line.startswith("###") or line.startswith("####")
        is_normal1 = next_line and not next_line.startswith("#") and not next_line.startswith("- ")
        is_normal2 = next_next_line and not next_next_line.startswith("#") and not next_next_line.startswith("- ")

        if is_heading and is_normal1 and is_normal2:
            title = re.sub(r"^#+", "", line).strip()
            institution = next_line.strip("*").strip()
            date = next_next_line.strip("*").strip()
            merged.append(f"**{title}**, *{institution}* · {date}")
            i += 3
        else:
            merged.append(line)
            i += 1

    return merged


# === Main Conversion Function ===

def convert_docx_to_md(input_path: str, output_path: str) -> None:
    
    doc = docx.Document(input_path)
    rels = extract_hyperlinks(doc)
    md_lines = []
    previous_style = ""

    for para in doc.paragraphs:
        line = process_paragraph(para, rels)
        if not line:
            continue

        current_style = para.style.name

        if previous_style in ("Bullet", "Normal") and current_style not in ("Bullet", "Normal"):
            md_lines.append("")

        md_lines.append(line)
        previous_style = current_style

    for table in doc.tables:
        md_lines.append(process_table(table))

    md_lines = merge_education_blocks(md_lines)
    detected_styles = set(para.style.name for para in doc.paragraphs if para.text.strip())
    print("🧾 Estilos detectados en el documento:")
    for s in sorted(detected_styles):
        print("  -", s)

    content = "\n".join(md_lines)
    content = normalize_spacing(content)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as md_file:
        md_file.write(content)

    print(f"✅ Markdown saved to {output_path}")
